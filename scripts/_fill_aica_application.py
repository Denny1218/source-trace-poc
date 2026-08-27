# -*- coding: utf-8 -*-
"""Fill AICA application docx from Pain Point row onward."""

from pathlib import Path

from docx import Document

SRC = Path(r"c:\Users\denny\Downloads\AICA 평가과제 신청서_유덕상.docx")
DST = Path(r"c:\Users\denny\Downloads\AICA 평가과제 신청서_유덕상_작성본.docx")
BAK = Path(r"c:\sourcechangeTrace\산출물\AICA_평가과제_신청서_유덕상_작성본.docx")
CHECK = Path(r"c:\sourcechangeTrace\_aica_filled_check.txt")


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    first.add_run(text)
    for p in list(cell.paragraphs[1:]):
        p._element.getparent().remove(p._element)


PAIN = (
    "장비 소스 유지보수 시 특정 함수나 코드 라인의 변경 시점·이유·관련 문서를 확인하려면 "
    "Git Commit/Diff와 변경내역서(PPT)를 각각 열어 수작업으로 대조해야 함. "
    "파일·함수 표기가 문서와 소스에 다르게 남는 경우가 있어 근거를 찾는 데 시간이 걸리고, "
    "경험이 적은 담당자는 관련 Commit이나 문서를 놓치기 쉬워 원인 파악이 지연되는 병목이 발생함."
)

SOLUTION = (
    "장비별 Git 이력과 변경내역서(PPT)를 Backend에 수집·연계하고, "
    "Web 및 IDE(VS Code / Eclipse / Visual Studio)에서 동일 API로 "
    "함수 단위 변경 이력과 선택 코드(라인) 변경 근거를 Markdown으로 조회하는 시스템(Source Trace)을 구축. "
    "1차 근거는 Git Diff·blame·line history이며, 규칙 기반 Evidence Link로 관련 변경내역서를 보강. "
    "설명 문장은 규칙 요약이 기본이고, 필요 시 Ollama로 표현을 다듬는 선택 경로를 둠 "
    "(벡터 임베딩 기반 RAG 없이 검색·규칙 중심의 근거 조회 구조)."
)

REPORT = (
    "본 과제는 유지보수 담당자가 함수·선택 코드의 변경 시점, Diff, 관련 변경내역서를 "
    "Web/IDE에서 같은 Backend 결과로 확인할 수 있게 하여, Git과 PPT를 따로 대조하던 조사 흐름을 단축하고 "
    "근거 확인의 일관성을 높임. "
    "서버 배포 패키지와 다중 IDE Adapter를 함께 정리해 사내 재현·인수인계·확장이 가능한 형태로 자산화하였으며, "
    "POC 범위에서 Backend API·Web·주요 IDE 연동까지 검증을 완료함."
)

GIT_LINK = (
    "(사내 Git 리포지토리 URL을 여기에 기입)\n"
    "예: https://gitlab.atec…/…/sourcechangeTrace\n"
    "※ 제출 소스 ZIP(SourceTrace_POC_Source.zip)과 동일 코드 기준으로 관리"
)

# 제출본 SourceTrace_POC_Source.zip 기준 (산출물/ 제외)
TREE = """sourcechangeTrace/
├── backend/
│   ├── app/
│   │   ├── api/
│   │   ├── core/
│   │   ├── db/
│   │   ├── schemas/
│   │   └── services/
│   └── tests/
├── frontend/
│   ├── public/
│   └── src/
│       ├── api/
│       ├── components/
│       ├── hooks/
│       ├── types/
│       └── utils/
├── vscode-extension/
│   ├── assets/
│   └── src/
├── eclipse-plugin/
│   ├── com.atec.sourcetrace.eclipse/
│   ├── feature/
│   └── update-site/
├── visualstudio-extension/
│   ├── vs2010/
│   └── vs2017/
├── scripts/
├── .env.example
├── OPERATING_TEST_STEP6.md
└── README.md"""


def main() -> None:
    doc = Document(str(SRC))
    table = doc.tables[0]
    contents = {
        8: PAIN,
        9: SOLUTION,
        10: REPORT,
        11: GIT_LINK,
        12: TREE,
    }
    for ri, text in contents.items():
        set_cell_text(table.rows[ri].cells[1], text)

    doc.save(str(DST))
    BAK.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(BAK))

    doc2 = Document(str(DST))
    t = doc2.tables[0]
    lines: list[str] = []
    for ri in range(8, 13):
        lines.append(f"=== R{ri} {t.rows[ri].cells[0].text.strip()[:50]} ===")
        lines.append(t.rows[ri].cells[1].text)
        lines.append("")
    CHECK.write_text("\n".join(lines), encoding="utf-8")
    print(f"saved={DST}")
    print(f"backup={BAK}")


if __name__ == "__main__":
    main()
